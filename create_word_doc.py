#!/usr/bin/env python3
"""
Generate Word Document for Email Campaign Manager Project
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os

def create_project_documentation():
    """Create comprehensive Word documentation for the project"""
    
    # Create new document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page
    title = doc.add_heading('Email Campaign Manager', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Email Marketing Solution')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.italic = True
    
    doc.add_paragraph()
    
    # Project details
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run('Project Documentation\n').bold = True
    details.add_run('Version: 1.0\n')
    details.add_run('Date: August 2024\n')
    details.add_run('Technology: Python, MongoDB, SMTP\n')
    
    # Page break
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Project Overview',
        '2. Features and Capabilities',
        '3. System Architecture',
        '4. Installation and Setup',
        '5. Configuration Guide',
        '6. User Guide',
        '7. API Documentation',
        '8. Database Schema',
        '9. Troubleshooting',
        '10. Security Considerations',
        '11. Performance Tips',
        '12. Future Enhancements'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 1. Project Overview
    doc.add_heading('1. Project Overview', level=1)
    
    overview_text = """
The Email Campaign Manager is a comprehensive Python application designed for managing email marketing campaigns with advanced features including contact management, template creation, list segmentation, and detailed analytics.

Key Highlights:
• Built with Python and MongoDB for scalability
• Gmail SMTP integration for reliable email delivery
• Environment-based configuration for security
• Comprehensive logging and analytics
• Personalization and template system
• CSV import/export capabilities
"""
    doc.add_paragraph(overview_text)
    
    # System Requirements
    doc.add_heading('System Requirements', level=2)
    req_table = doc.add_table(rows=1, cols=2)
    req_table.style = 'Table Grid'
    req_hdr_cells = req_table.rows[0].cells
    req_hdr_cells[0].text = 'Component'
    req_hdr_cells[1].text = 'Requirement'
    
    requirements = [
        ('Python Version', 'Python 3.7 or higher'),
        ('Database', 'MongoDB (Local or Atlas)'),
        ('Email Service', 'SMTP Server (Gmail supported)'),
        ('Dependencies', 'pymongo, pandas, python-dotenv'),
        ('Operating System', 'Windows, macOS, Linux'),
        ('Memory', 'Minimum 512MB RAM'),
        ('Storage', 'Minimum 100MB free space')
    ]
    
    for req, desc in requirements:
        row_cells = req_table.add_row().cells
        row_cells[0].text = req
        row_cells[1].text = desc
    
    doc.add_page_break()
    
    # 2. Features and Capabilities
    doc.add_heading('2. Features and Capabilities', level=1)
    
    # Contact Management
    doc.add_heading('Contact Management', level=2)
    contact_features = [
        'Add, update, and delete contacts with validation',
        'Custom fields support for additional contact information',
        'CSV import with flexible field mapping',
        'Contact export functionality',
        'Duplicate prevention and email validation',
        'Contact search and filtering capabilities'
    ]
    
    for feature in contact_features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    # Email Templates
    doc.add_heading('Email Templates', level=2)
    template_features = [
        'Create reusable email templates',
        'Support for both plain text and HTML content',
        'Dynamic content personalization with placeholders',
        'Template versioning and management',
        'Rich text formatting options'
    ]
    
    for feature in template_features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    # Campaign Management
    doc.add_heading('Campaign Management', level=2)
    campaign_features = [
        'Create targeted email campaigns',
        'Multi-list targeting capabilities',
        'Real-time campaign statistics',
        'Campaign status tracking (draft, sent)',
        'Delivery success/failure reporting',
        'Campaign scheduling and automation'
    ]
    
    for feature in campaign_features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # 3. System Architecture
    doc.add_heading('3. System Architecture', level=1)
    
    arch_text = """
The Email Campaign Manager follows a modular architecture with clear separation of concerns:

Application Layer:
• EmailCampaignManager: Main controller class
• Contact, EmailTemplate, Campaign: Data models
• SMTP Handler: Email delivery management

Data Layer:
• MongoDB Collections: contacts, templates, campaigns, contact_lists, email_logs
• Indexed fields for performance optimization
• Automatic data validation and constraints

Configuration Layer:
• Environment-based configuration (.env)
• Secure credential management
• Flexible SMTP settings

Integration Layer:
• Gmail SMTP integration
• CSV import/export functionality
• Logging and monitoring system
"""
    doc.add_paragraph(arch_text)
    
    # Database Collections
    doc.add_heading('Database Collections', level=2)
    collections_table = doc.add_table(rows=1, cols=3)
    collections_table.style = 'Table Grid'
    hdr_cells = collections_table.rows[0].cells
    hdr_cells[0].text = 'Collection'
    hdr_cells[1].text = 'Purpose'
    hdr_cells[2].text = 'Key Fields'
    
    collections = [
        ('contacts', 'Store contact information', 'email, first_name, last_name, company, custom_fields'),
        ('templates', 'Email templates', 'name, subject, body, html_body, created_at'),
        ('campaigns', 'Campaign definitions', 'name, template_id, contact_list_ids, status, statistics'),
        ('contact_lists', 'Contact list organization', 'name, description, contact_ids'),
        ('email_logs', 'Delivery tracking', 'to_email, subject, status, timestamp, error_message')
    ]
    
    for coll, purpose, fields in collections:
        row_cells = collections_table.add_row().cells
        row_cells[0].text = coll
        row_cells[1].text = purpose
        row_cells[2].text = fields
    
    doc.add_page_break()
    
    # 4. Installation and Setup
    doc.add_heading('4. Installation and Setup', level=1)
    
    install_text = """
Follow these steps to set up the Email Campaign Manager:
"""
    doc.add_paragraph(install_text)
    
    # Installation Steps
    steps = [
        ('Clone Repository', 'git clone <repository-url>\ncd Email_Camp'),
        ('Install Dependencies', 'pip install pymongo pandas python-dotenv'),
        ('Setup MongoDB', 'Install MongoDB locally or create MongoDB Atlas cluster'),
        ('Configure Environment', 'Copy .env.example to .env and update settings'),
        ('Test Setup', 'python test_setup.py'),
        ('Run Application', 'python mail.py')
    ]
    
    for i, (step, command) in enumerate(steps, 1):
        doc.add_heading(f'Step {i}: {step}', level=3)
        if '\\n' in command:
            for cmd in command.split('\\n'):
                p = doc.add_paragraph(cmd, style='Intense Quote')
        else:
            doc.add_paragraph(command, style='Intense Quote')
    
    doc.add_page_break()
    
    # 5. Configuration Guide
    doc.add_heading('5. Configuration Guide', level=1)
    
    config_text = """
The application uses environment variables for configuration. All settings are stored in the .env file for security and flexibility.
"""
    doc.add_paragraph(config_text)
    
    # Configuration Table
    config_table = doc.add_table(rows=1, cols=3)
    config_table.style = 'Table Grid'
    config_hdr_cells = config_table.rows[0].cells
    config_hdr_cells[0].text = 'Variable'
    config_hdr_cells[1].text = 'Description'
    config_hdr_cells[2].text = 'Example'
    
    config_vars = [
        ('MONGODB_URI', 'MongoDB connection string', 'mongodb://localhost:27017/'),
        ('MONGODB_DATABASE', 'Database name', 'email_campaigns'),
        ('SMTP_SERVER', 'SMTP server hostname', 'smtp.gmail.com'),
        ('SMTP_PORT', 'SMTP server port', '587'),
        ('SMTP_USERNAME', 'SMTP username', 'user@gmail.com'),
        ('SMTP_PASSWORD', 'SMTP password/app password', 'app_password_16_chars'),
        ('SMTP_USE_TLS', 'Enable TLS encryption', 'True'),
        ('FROM_EMAIL', 'Sender email address', 'sender@gmail.com'),
        ('FROM_NAME', 'Sender display name', 'Your Name'),
        ('RECEIVER_EMAILS', 'Recipient emails (comma-separated)', 'user1@example.com,user2@example.com'),
        ('DEFAULT_CONTACT_LIST_NAME', 'Default contact list name', 'Default Recipients')
    ]
    
    for var, desc, example in config_vars:
        row_cells = config_table.add_row().cells
        row_cells[0].text = var
        row_cells[1].text = desc
        row_cells[2].text = example
    
    doc.add_page_break()
    
    # 6. User Guide
    doc.add_heading('6. User Guide', level=1)
    
    # Quick Start
    doc.add_heading('Quick Start Example', level=2)
    quickstart_code = '''
from mail import EmailCampaignManager, Contact, EmailTemplate, Campaign

# Initialize manager
manager = EmailCampaignManager()

# Add contact
contact = Contact(
    email="user@example.com",
    first_name="John",
    last_name="Doe",
    company="Example Corp"
)
manager.add_contact(contact)

# Create template
template = EmailTemplate(
    name="Welcome Email",
    subject="Welcome {{first_name}}!",
    body="Hello {{first_name}}, welcome to our service!"
)
template_id = manager.save_template(template)

# Create campaign
campaign = Campaign(
    name="Welcome Campaign",
    template_id=template_id,
    contact_list_ids=[list_id]
)
campaign_id = manager.create_campaign(campaign)

# Send campaign
stats = manager.send_campaign("Welcome Campaign")
print(f"Campaign results: {stats}")
'''
    doc.add_paragraph(quickstart_code, style='Intense Quote')
    
    # Personalization Guide
    doc.add_heading('Email Personalization', level=2)
    personalization_text = """
Use these placeholders in your email templates for dynamic content:

• {{first_name}} - Contact's first name
• {{last_name}} - Contact's last name
• {{full_name}} - Full name (first + last)
• {{email}} - Contact's email address
• {{company}} - Contact's company
• {{custom.field_name}} - Custom field values

Example Template:
Subject: Welcome {{first_name}}!

Hello {{first_name}} {{last_name}},

Thank you for joining us at {{company}}.
Your department: {{custom.department}}

Best regards,
The Team
"""
    doc.add_paragraph(personalization_text)
    
    doc.add_page_break()
    
    # 7. API Documentation
    doc.add_heading('7. API Documentation', level=1)
    
    api_text = """
The EmailCampaignManager class provides the main API for interacting with the system.
"""
    doc.add_paragraph(api_text)
    
    # Key Methods
    methods = [
        ('Contact Management', [
            'add_contact(contact: Contact) -> str',
            'get_contact(email: str) -> Optional[Contact]',
            'update_contact(email: str, updates: Dict) -> bool',
            'delete_contact(email: str) -> bool',
            'import_contacts_from_csv(file_path: str, mapping: Dict) -> int'
        ]),
        ('Template Management', [
            'save_template(template: EmailTemplate) -> str',
            'get_template(name: str) -> Optional[EmailTemplate]',
            'list_templates() -> List[str]',
            'delete_template(name: str) -> bool'
        ]),
        ('Campaign Management', [
            'create_campaign(campaign: Campaign) -> str',
            'send_campaign(name: str, attachments: List[str]) -> Dict',
            'get_campaign_stats(name: str) -> Dict',
            'list_campaigns() -> List[str]'
        ]),
        ('List Management', [
            'create_contact_list(name: str, description: str) -> str',
            'add_contacts_to_list(list_name: str, emails: List[str]) -> int',
            'get_contact_list_contacts(list_name: str) -> List[Contact]'
        ])
    ]
    
    for category, method_list in methods:
        doc.add_heading(category, level=2)
        for method in method_list:
            doc.add_paragraph(method, style='List Bullet')
    
    doc.add_page_break()
    
    # 8. Security Considerations
    doc.add_heading('8. Security Considerations', level=1)
    
    security_text = """
Important security practices for production deployment:

Authentication and Credentials:
• Use Gmail App Passwords instead of regular passwords
• Store sensitive credentials in .env file (never commit to version control)
• Implement proper access controls for the application
• Regularly rotate SMTP passwords

Data Protection:
• Validate all email addresses before storing
• Implement rate limiting for email sending
• Monitor email logs for suspicious activity
• Encrypt sensitive data in the database

Network Security:
• Use TLS encryption for SMTP connections
• Implement firewall rules for MongoDB access
• Use secure MongoDB connection strings
• Consider VPN access for production databases

Application Security:
• Implement input validation and sanitization
• Use parameterized queries to prevent injection attacks
• Implement proper error handling without exposing sensitive information
• Regular security audits and updates
"""
    doc.add_paragraph(security_text)
    
    doc.add_page_break()
    
    # 9. Troubleshooting
    doc.add_heading('9. Troubleshooting', level=1)
    
    # Common Issues Table
    issues_table = doc.add_table(rows=1, cols=3)
    issues_table.style = 'Table Grid'
    issues_hdr_cells = issues_table.rows[0].cells
    issues_hdr_cells[0].text = 'Issue'
    issues_hdr_cells[1].text = 'Symptoms'
    issues_hdr_cells[2].text = 'Solution'
    
    issues = [
        ('SMTP Authentication Failed', 'Error: Authentication failed', 'Check Gmail App Password, enable 2FA'),
        ('MongoDB Connection Error', 'Error: Failed to connect to MongoDB', 'Verify MongoDB URI, check network connectivity'),
        ('Email Not Delivered', 'Status: failed in logs', 'Check recipient email, review SMTP settings'),
        ('Connection Unexpectedly Closed', 'SMTP connection drops', 'Use proper EHLO hostname (gmail.com)'),
        ('Invalid Email Address', 'Validation errors', 'Check email format, use proper regex validation'),
        ('Template Not Found', 'Template errors in campaigns', 'Verify template name, check database'),
        ('Permission Denied', 'Database write errors', 'Check MongoDB permissions, connection string'),
        ('Import CSV Failed', 'CSV import errors', 'Verify file format, check field mapping')
    ]
    
    for issue, symptom, solution in issues:
        row_cells = issues_table.add_row().cells
        row_cells[0].text = issue
        row_cells[1].text = symptom
        row_cells[2].text = solution
    
    doc.add_page_break()
    
    # 10. Performance Tips
    doc.add_heading('10. Performance Tips', level=1)
    
    performance_text = """
Optimization strategies for large-scale deployments:

Database Optimization:
• Use indexes on frequently queried fields (email, name)
• Implement database connection pooling
• Regular cleanup of old email logs
• Consider sharding for very large datasets

Email Delivery Optimization:
• Implement batch processing for large campaigns
• Use connection pooling for SMTP
• Implement retry logic with exponential backoff
• Monitor and respect rate limits

Application Performance:
• Use asynchronous processing for large operations
• Implement caching for frequently accessed data
• Optimize database queries to reduce round trips
• Monitor memory usage and implement pagination

Monitoring and Maintenance:
• Implement comprehensive logging
• Set up alerts for failed email deliveries
• Monitor database performance metrics
• Regular backup and disaster recovery procedures
"""
    doc.add_paragraph(performance_text)
    
    doc.add_page_break()
    
    # 11. Future Enhancements
    doc.add_heading('11. Future Enhancements', level=1)
    
    enhancements = [
        'Web-based user interface with dashboard',
        'Advanced analytics and reporting features',
        'A/B testing capabilities for campaigns',
        'Email automation and drip campaigns',
        'Integration with CRM systems',
        'Advanced segmentation and targeting',
        'Email template builder with drag-and-drop',
        'Multi-language support and localization',
        'Integration with social media platforms',
        'Advanced spam filtering and compliance tools',
        'Mobile application for campaign management',
        'API endpoints for third-party integrations'
    ]
    
    for enhancement in enhancements:
        doc.add_paragraph(enhancement, style='List Bullet')
    
    doc.add_page_break()
    
    # 12. Appendix
    doc.add_heading('12. Appendix', level=1)
    
    doc.add_heading('File Structure', level=2)
    file_structure = """
Email_Camp/
├── mail.py                 # Main application file
├── .env                    # Environment configuration
├── .env.example           # Configuration template
├── requirements.txt       # Python dependencies
├── .gitignore            # Git ignore rules
├── README.md             # Project documentation
├── API_DOCUMENTATION.md  # Detailed API docs
├── test_setup.py         # Setup verification script
├── smtp_test.py          # SMTP connection testing
└── create_word_doc.py    # Documentation generator
"""
    doc.add_paragraph(file_structure, style='Intense Quote')
    
    doc.add_heading('Environment Variables Reference', level=2)
    env_example = """
# MongoDB Configuration
MONGODB_URI=mongodb://localhost:27017/
MONGODB_DATABASE=email_campaigns

# SMTP Configuration
SMTP_SERVER=smtp.gmail.com
SMTP_PORT=587
SMTP_USERNAME=your_email@gmail.com
SMTP_PASSWORD=your_app_password
SMTP_USE_TLS=True

# Email Settings
FROM_EMAIL=your_email@gmail.com
FROM_NAME=Your Name

# Receiver Configuration
RECEIVER_EMAILS=user1@example.com,user2@example.com
DEFAULT_CONTACT_LIST_NAME=Default Recipients
"""
    doc.add_paragraph(env_example, style='Intense Quote')
    
    # Save document
    doc_path = 'Email_Campaign_Manager_Documentation.docx'
    doc.save(doc_path)
    print(f"Documentation generated: {doc_path}")
    
    return doc_path

if __name__ == "__main__":
    try:
        doc_path = create_project_documentation()
        print(f"[OK] Word document created successfully: {doc_path}")
    except Exception as e:
        print(f"Error creating document: {e}")
        # Install python-docx if not available
        if "No module named 'docx'" in str(e):
            print("Installing python-docx...")
            import subprocess
            subprocess.run(['pip', 'install', 'python-docx'])
            print("Please run the script again after installation.")
        else:
            import traceback
            traceback.print_exc()